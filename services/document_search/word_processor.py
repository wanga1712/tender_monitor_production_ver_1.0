"""
MODULE: services.document_search.word_processor
RESPONSIBILITY: Extract text from Word documents (.docx, .doc).
ALLOWED: docx, win32com, subprocess, logging, core.exceptions, file_lock_handler.
FORBIDDEN: Direct database access.
ERRORS: DocumentSearchError.

Модуль для обработки Word документов (.docx, .doc).

Поддерживает:
- .docx файлы - через python-docx
- .doc файлы (старый формат) - через антивирус/word (или fallback на текст)
"""

from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, Iterable, Optional
import io
import os
import subprocess

from loguru import logger

from core.exceptions import DocumentSearchError
from services.document_search.file_lock_handler import handle_file_lock


class WordProcessor:
    """Класс для обработки Word документов."""
    
    def __init__(self):
        """Инициализация процессора Word."""
        self._check_dependencies()
    
    def _check_dependencies(self) -> None:
        """Проверка наличия необходимых библиотек."""
        try:
            import docx
            self._docx_available = True
        except ImportError:
            self._docx_available = False
            logger.debug("python-docx не установлен, обработка .docx файлов будет недоступна")
        
        # Для старых .doc файлов пробуем антивирус или другие методы
        self._doc_available = False
        try:
            # Пробуем использовать antiword (если доступен)
            import subprocess
            result = subprocess.run(
                ['antiword', '-h'],
                capture_output=True,
                timeout=2,
                encoding='utf-8',
                errors='replace',  # Заменяем некорректные символы вместо ошибки
            )
            if result.returncode == 0:
                self._doc_available = True
                self._doc_method = 'antiword'
                logger.debug("antiword доступен для обработки .doc файлов")
        except (FileNotFoundError, subprocess.TimeoutExpired, ImportError):
            pass
        
        # Если antiword не доступен, пробуем другие методы
        if not self._doc_available:
            try:
                # В Windows можно использовать Word COM объект (если установлен Word)
                import win32com.client
                self._doc_available = True
                self._doc_method = 'com'
                logger.debug("Word COM объект доступен для обработки .doc файлов")
            except ImportError:
                self._doc_available = False
                logger.debug(
                    "Обработка .doc файлов будет недоступна. "
                    "Установите python-docx для .docx или Word COM для .doc"
                )
    
    def is_word_file(self, file_path: Path) -> bool:
        """Проверка, является ли файл Word документом."""
        if not file_path.exists():
            return False
        suffix = file_path.suffix.lower()
        return suffix in ('.docx', '.doc')
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """
        Извлечение текста из .docx файла.
        
        Args:
            file_path: Путь к .docx файлу
            
        Returns:
            Извлеченный текст
            
        Raises:
            DocumentSearchError: Если не удалось извлечь текст
        """
        if not self._docx_available:
            raise DocumentSearchError(
                "python-docx не установлен. Установите: pip install python-docx"
            )
        
        if not self.is_word_file(file_path) or file_path.suffix.lower() != '.docx':
            raise DocumentSearchError(f"Файл {file_path} не является .docx файлом")
        
        try:
            import docx
            
            # Используем handle_file_lock для обработки блокировок
            doc = handle_file_lock(file_path, lambda: docx.Document(file_path))
            text_parts = []
            
            # Извлекаем текст из всех параграфов
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            if not text_parts:
                raise DocumentSearchError(f"Документ {file_path.name} не содержит текста")
            
            full_text = "\n".join(text_parts)
            logger.info(f"Извлечен текст из Word {file_path.name}: {len(full_text)} символов")
            return full_text
            
        except ImportError as import_error:
            raise DocumentSearchError(f"python-docx не установлен: {import_error}")
        except Exception as error:
            raise DocumentSearchError(f"Не удалось извлечь текст из Word {file_path.name}: {error}")
    
    def extract_text_from_doc(self, file_path: Path) -> str:
        """
        Извлечение текста из старого .doc файла.
        
        Args:
            file_path: Путь к .doc файлу
            
        Returns:
            Извлеченный текст
            
        Raises:
            DocumentSearchError: Если не удалось извлечь текст
        """
        if not self._doc_available:
            raise DocumentSearchError(
                "Обработка .doc файлов недоступна. Установите antiword или используйте Word COM объект"
            )
        
        if not self.is_word_file(file_path) or file_path.suffix.lower() != '.doc':
            raise DocumentSearchError(f"Файл {file_path} не является .doc файлом")
        
        try:
            if self._doc_method == 'antiword':
                import subprocess
                
                result = subprocess.run(
                    ['antiword', str(file_path)],
                    capture_output=True,
                    text=True,
                    timeout=30,
                    encoding='utf-8',
                    errors='replace',  # Заменяем некорректные символы вместо ошибки
                )
                
                if result.returncode != 0:
                    raise DocumentSearchError(f"antiword не смог обработать {file_path.name}: {result.stderr}")
                
                if not result.stdout.strip():
                    raise DocumentSearchError(f"Документ {file_path.name} не содержит текста")
                
                logger.info(f"Извлечен текст из Word .doc {file_path.name} через antiword: {len(result.stdout)} символов")
                return result.stdout
                
            elif self._doc_method == 'com':
                # Используем Word COM объект (Windows)
                import win32com.client
                import pythoncom
                
                pythoncom.CoInitialize()
                
                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    
                    doc = word.Documents.Open(str(file_path.absolute()))
                    text = doc.Content.Text
                    doc.Close()
                    word.Quit()
                    
                    if not text.strip():
                        raise DocumentSearchError(f"Документ {file_path.name} не содержит текста")
                    
                    logger.info(f"Извлечен текст из Word .doc {file_path.name} через COM: {len(text)} символов")
                    return text
                    
                finally:
                    pythoncom.CoUninitialize()
            else:
                raise DocumentSearchError(f"Неизвестный метод обработки .doc: {self._doc_method}")
                
        except ImportError as import_error:
            raise DocumentSearchError(f"Необходимые библиотеки не установлены: {import_error}")
        except subprocess.TimeoutExpired:
            raise DocumentSearchError(f"Превышено время ожидания при обработке {file_path.name}")
        except Exception as error:
            raise DocumentSearchError(f"Не удалось извлечь текст из Word .doc {file_path.name}: {error}")
    
    def extract_text(self, file_path: Path) -> str:
        """
        Универсальный метод извлечения текста из Word документа.
        
        Автоматически определяет формат (.docx или .doc) и использует соответствующий метод.
        
        Args:
            file_path: Путь к Word файлу
            
        Returns:
            Извлеченный текст
            
        Raises:
            DocumentSearchError: Если не удалось извлечь текст
        """
        if not self.is_word_file(file_path):
            raise DocumentSearchError(f"Файл {file_path} не является Word документом")
        
        suffix = file_path.suffix.lower()
        
        if suffix == '.docx':
            return self.extract_text_from_docx(file_path)
        elif suffix == '.doc':
            return self.extract_text_from_doc(file_path)
        else:
            raise DocumentSearchError(f"Неподдерживаемый формат Word документа: {suffix}")
    
    def iter_word_cells(self, file_path: Path) -> Iterable[Dict[str, Any]]:
        """
        Итератор по тексту Word документа, имитирующий структуру Excel ячеек.
        
        Разбивает текст на "ячейки" (слова/фразы) для совместимости с поиском.
        
        Args:
            file_path: Путь к Word файлу
            
        Yields:
            Dict с полями:
            - text: извлеченный текст
            - display_text: текст для отображения
            - sheet_name: имя "листа" (используем имя файла)
            - row: номер строки
            - column: номер колонки
            - cell_address: адрес ячейки
        """
        try:
            full_text = self.extract_text(file_path)
        except DocumentSearchError as error:
            logger.error(f"Не удалось извлечь текст из Word {file_path}: {error}")
            return
        
        # Разбиваем текст на строки и слова для поиска
        lines = full_text.split('\n')
        sheet_name = file_path.stem
        
        for row_num, line in enumerate(lines, 1):
            if not line.strip():
                continue
            
            # Разбиваем строку на слова/фразы
            words = line.split()
            for col_num, word in enumerate(words, 1):
                yield {
                    "text": word.lower(),
                    "display_text": word,
                    "sheet_name": sheet_name,
                    "row": row_num,
                    "column": col_num,
                    "cell_address": f"{row_num}:{col_num}",
                }

